"""
File handling utilities for various document formats.
"""

import os
from typing import Union, Dict, Any
import logging
from pathlib import Path


class FileHandler:
    """Handle reading and writing of various file formats."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.supported_formats = {'.txt', '.md', '.docx', '.doc'}
    
    def read_file(self, file_path: str) -> str:
        """
        Read file content based on file extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            File content as string
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: {self.supported_formats}")
        
        try:
            if file_ext in ['.txt', '.md']:
                return self._read_text_file(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._read_docx_file(file_path)
            else:
                raise ValueError(f"Handler not implemented for {file_ext}")
                
        except Exception as e:
            self.logger.error(f"Error reading file {file_path}: {e}")
            raise
    
    def write_file(self, file_path: str, content: str, format_hint: str = None) -> None:
        """
        Write content to file based on extension or format hint.
        
        Args:
            file_path: Output file path
            content: Content to write
            format_hint: Optional format hint if extension is ambiguous
        """
        try:
            # Create directory if it doesn't exist
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext in ['.txt', '.md'] or not file_ext:
                self._write_text_file(file_path, content)
            elif file_ext in ['.docx']:
                self._write_docx_file(file_path, content)
            else:
                # Default to text file
                self._write_text_file(file_path, content)
                
            self.logger.info(f"Successfully wrote file: {file_path}")
            
        except Exception as e:
            self.logger.error(f"Error writing file {file_path}: {e}")
            raise
    
    def _read_text_file(self, file_path: str) -> str:
        """Read plain text or markdown file."""
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                self.logger.debug(f"Successfully read {file_path} with {encoding} encoding")
                return content
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode file {file_path} with any supported encoding")
    
    def _write_text_file(self, file_path: str, content: str) -> None:
        """Write plain text or markdown file."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    def _read_docx_file(self, file_path: str) -> str:
        """Read Microsoft Word document."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx package required for .docx files. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            return '\n\n'.join(paragraphs)
            
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {e}")
    
    def _write_docx_file(self, file_path: str, content: str) -> None:
        """Write Microsoft Word document."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx package required for .docx files. Install with: pip install python-docx")
        
        try:
            doc = Document()
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it's a header (starts with #)
                    if para_text.strip().startswith('#'):
                        level = len(para_text) - len(para_text.lstrip('#'))
                        header_text = para_text.lstrip('#').strip()
                        doc.add_heading(header_text, level=min(level, 9))
                    else:
                        doc.add_paragraph(para_text.strip())
            
            doc.save(file_path)
            
        except Exception as e:
            raise ValueError(f"Error writing DOCX file: {e}")
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get information about a file."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        stat = os.stat(file_path)
        file_ext = Path(file_path).suffix.lower()
        
        # Get word count estimate
        try:
            content = self.read_file(file_path)
            word_count = len(content.split())
            char_count = len(content)
        except:
            word_count = None
            char_count = None
        
        return {
            "path": file_path,
            "size_bytes": stat.st_size,
            "extension": file_ext,
            "supported": file_ext in self.supported_formats,
            "word_count": word_count,
            "char_count": char_count,
            "modified_time": stat.st_mtime
        }
    
    def validate_file(self, file_path: str) -> Dict[str, Any]:
        """Validate if file can be processed."""
        validation_result = {
            "valid": False,
            "errors": [],
            "warnings": [],
            "info": {}
        }
        
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                validation_result["errors"].append("File does not exist")
                return validation_result
            
            # Get file info
            file_info = self.get_file_info(file_path)
            validation_result["info"] = file_info
            
            # Check format support
            if not file_info["supported"]:
                validation_result["errors"].append(f"Unsupported format: {file_info['extension']}")
                return validation_result
            
            # Check file size (warn if > 10MB)
            if file_info["size_bytes"] > 10 * 1024 * 1024:
                validation_result["warnings"].append("Large file size may result in high API costs")
            
            # Try to read file
            try:
                content = self.read_file(file_path)
                if not content.strip():
                    validation_result["errors"].append("File is empty")
                    return validation_result
                
                # Warn about very long documents
                if file_info["word_count"] and file_info["word_count"] > 50000:
                    validation_result["warnings"].append("Very long document may take significant time to process")
                
            except Exception as e:
                validation_result["errors"].append(f"Cannot read file: {str(e)}")
                return validation_result
            
            validation_result["valid"] = True
            
        except Exception as e:
            validation_result["errors"].append(f"Validation error: {str(e)}")
        
        return validation_result